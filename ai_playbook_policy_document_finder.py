#!/usr/bin/env python3
"""
AI Playbook Public Policy Document Finder

Finds and organizes authoritative public policy documents for an AI Playbook
workflow. The tool uses the official Google Custom Search JSON API when
credentials are available. If credentials are missing, it falls back to a manual
seed URL file so the script can still produce useful assignment artifacts.

Run:
    python ai_playbook_policy_document_finder.py --quick

Outputs:
    runs/run_YYYYMMDD_HHMMSS_microseconds/policy_document_candidates_RUNID.csv
    runs/run_YYYYMMDD_HHMMSS_microseconds/policy_document_report_RUNID.md
    runs/run_YYYYMMDD_HHMMSS_microseconds/policy_document_report_RUNID.docx
    runs/run_YYYYMMDD_HHMMSS_microseconds/run_log_RUNID.txt
"""

from __future__ import annotations

import argparse
import csv
import html
import json
import os
import re
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
import urllib.robotparser
from dataclasses import dataclass
from datetime import date, datetime
from typing import Any, Dict, Iterable, List, Optional, Tuple

try:
    import requests  # type: ignore
except ImportError:  # pragma: no cover - fallback keeps the script dependency-free.
    requests = None


GOOGLE_CUSTOM_SEARCH_ENDPOINT = "https://www.googleapis.com/customsearch/v1"
RUNS_DIR = "runs"
RUN_INDEX_CSV = os.path.join(RUNS_DIR, "run_index.csv")
LATEST_OUTPUT_CSV = "policy_document_candidates.csv"
LATEST_OUTPUT_MD = "policy_document_report.md"
LATEST_OUTPUT_DOCX = "policy_document_report.docx"
DEFAULT_SEED_FILE = "seed_urls.txt"
DEFAULT_MAX_RESULTS = 20
QUICK_MAX_RESULTS = 3
DEFAULT_SLEEP_SECONDS = 0.75
QUICK_SLEEP_SECONDS = 0.2
DEFAULT_TIMEOUT_SECONDS = 10.0
QUICK_TIMEOUT_SECONDS = 5.0
USER_AGENT = (
    "AIPlaybookPolicyDocumentFinder/1.0 "
    "(public administration coursework; respectful metadata collection)"
)


GOOGLE_POLICY_QUERIES = [
    "site:.gov public policy report filetype:pdf",
    "site:gao.gov public policy report",
    "site:congress.gov public policy report",
    "site:whitehouse.gov policy report",
    "site:cbo.gov policy report",
    "site:urban.org public policy report",
    "site:brookings.edu public policy report",
    "site:rand.org public policy report",
    "site:pewresearch.org public policy report",
]


CSV_COLUMNS = [
    "title",
    "source_organization",
    "url",
    "domain",
    "file_type",
    "publication_year",
    "policy_topic",
    "document_type",
    "authority_level",
    "relevance_note",
    "retrieval_date",
    "source_mode",
]


RUN_INDEX_COLUMNS = [
    "run_id",
    "datetime",
    "quick_mode",
    "source_mode",
    "seed_file",
    "topic",
    "max_results_requested",
    "documents_collected",
    "run_folder",
    "csv_output_path",
    "markdown_report_path",
    "word_report_path",
    "log_path",
    "also_write_latest",
]


SOURCE_ORGANIZATIONS = {
    "gao.gov": "U.S. Government Accountability Office",
    "congress.gov": "Congress.gov / Congressional Research Service",
    "crsreports.congress.gov": "Congressional Research Service",
    "cbo.gov": "Congressional Budget Office",
    "whitehouse.gov": "The White House",
    "govinfo.gov": "U.S. Government Publishing Office",
    "ed.gov": "U.S. Department of Education",
    "hhs.gov": "U.S. Department of Health and Human Services",
    "hud.gov": "U.S. Department of Housing and Urban Development",
    "epa.gov": "U.S. Environmental Protection Agency",
    "transportation.gov": "U.S. Department of Transportation",
    "urban.org": "Urban Institute",
    "brookings.edu": "Brookings Institution",
    "rand.org": "RAND Corporation",
    "pewresearch.org": "Pew Research Center",
    "kff.org": "KFF",
    "cbpp.org": "Center on Budget and Policy Priorities",
}


ESTABLISHED_THINK_TANK_DOMAINS = {
    "urban.org",
    "brookings.edu",
    "pewresearch.org",
    "kff.org",
    "cbpp.org",
}


RESEARCH_INSTITUTION_DOMAINS = {
    "rand.org",
}


CONGRESSIONAL_OVERSIGHT_DOMAINS = {
    "gao.gov",
    "cbo.gov",
    "congress.gov",
    "crsreports.congress.gov",
}


POLICY_TOPIC_KEYWORDS = {
    "housing policy": [
        "housing",
        "rental",
        "rent",
        "homeless",
        "affordable housing",
        "zoning",
        "mortgage",
    ],
    "health policy": [
        "health",
        "healthcare",
        "medicaid",
        "medicare",
        "public health",
        "insurance",
    ],
    "education policy": [
        "education",
        "school",
        "student",
        "teacher",
        "higher education",
        "k-12",
    ],
    "budget and fiscal policy": [
        "budget",
        "fiscal",
        "tax",
        "deficit",
        "appropriation",
        "spending",
        "economic outlook",
    ],
    "technology and AI policy": [
        "artificial intelligence",
        " ai ",
        "algorithm",
        "automation",
        "cybersecurity",
        "data privacy",
        "technology",
    ],
    "climate and environment policy": [
        "climate",
        "environment",
        "emissions",
        "energy",
        "resilience",
        "pollution",
    ],
    "labor and workforce policy": [
        "labor",
        "workforce",
        "employment",
        "worker",
        "jobs",
        "wage",
    ],
    "criminal justice policy": [
        "criminal justice",
        "policing",
        "prison",
        "incarceration",
        "courts",
    ],
    "immigration policy": [
        "immigration",
        "immigrant",
        "border",
        "asylum",
        "visa",
    ],
    "governance and public administration": [
        "governance",
        "public administration",
        "government performance",
        "federal agency",
        "oversight",
        "implementation",
    ],
}


@dataclass
class PolicyDocument:
    title: str
    source_organization: str
    url: str
    domain: str
    file_type: str
    publication_year: str
    policy_topic: str
    document_type: str
    authority_level: str
    relevance_note: str
    retrieval_date: str
    source_mode: str
    snippet: str = ""
    fetch_status: str = "Not fetched"


@dataclass
class HttpResult:
    status_code: int
    headers: Dict[str, str]
    body: bytes
    final_url: str
    error: str = ""


@dataclass
class RunPaths:
    run_id: str
    run_folder: str
    csv_path: str
    report_path: str
    docx_path: str
    log_path: str


class RobotsCache:
    """Fetch and cache robots.txt rules before page content is requested."""

    def __init__(self, sleep_seconds: float, timeout_seconds: float, quick: bool) -> None:
        self.sleep_seconds = sleep_seconds
        self.timeout_seconds = timeout_seconds
        self.quick = quick
        self.cache: Dict[str, Optional[urllib.robotparser.RobotFileParser]] = {}

    def can_fetch(self, url: str) -> Tuple[bool, str]:
        parsed = urllib.parse.urlparse(url)
        if not parsed.scheme or not parsed.netloc:
            return False, "Invalid URL"

        domain = parsed.netloc.lower()
        if domain not in self.cache:
            self.cache[domain] = self._load_parser(parsed)

        parser = self.cache[domain]
        if parser is None:
            return True, "robots.txt unavailable; proceeding cautiously"
        allowed = parser.can_fetch(USER_AGENT, url)
        if allowed:
            return True, "Allowed by robots.txt"
        return False, "Blocked by robots.txt"

    def _load_parser(
        self, parsed_url: urllib.parse.ParseResult
    ) -> Optional[urllib.robotparser.RobotFileParser]:
        robots_url = f"{parsed_url.scheme}://{parsed_url.netloc}/robots.txt"
        result = http_request(
            robots_url,
            max_bytes=250_000,
            max_retries=1,
            sleep_seconds=self.sleep_seconds,
            timeout_seconds=self.timeout_seconds,
        )
        if result.status_code < 200 or result.status_code >= 400 or not result.body:
            return None

        parser = urllib.robotparser.RobotFileParser()
        parser.set_url(robots_url)
        robots_text = decode_body(result.body, result.headers)
        parser.parse(robots_text.splitlines())
        return parser


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = html.unescape(str(value))
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def keyword_hits(text: str, keywords: Iterable[str]) -> List[str]:
    padded_text = f" {text.lower()} "
    hits = []
    for keyword in keywords:
        lowered = keyword.lower()
        if lowered.strip() == "ai":
            pattern = r"\bai\b"
        else:
            pattern = re.escape(lowered).replace(r"\ ", r"\s+")
        if re.search(pattern, padded_text):
            hits.append(keyword.strip())
    return hits


def domain_from_url(url: str) -> str:
    parsed = urllib.parse.urlparse(url)
    domain = parsed.netloc.lower().strip()
    if domain.startswith("www."):
        domain = domain[4:]
    return domain


def base_domain_match(domain: str, known_domain: str) -> bool:
    return domain == known_domain or domain.endswith(f".{known_domain}")


def source_organization_for_domain(domain: str) -> str:
    for known_domain, organization in SOURCE_ORGANIZATIONS.items():
        if base_domain_match(domain, known_domain):
            return organization
    if domain.endswith(".gov"):
        return "Official government source"
    if domain.endswith(".edu"):
        return "University or research institution"
    return domain or "Unknown source"


def authority_level_for_domain(domain: str) -> str:
    if any(base_domain_match(domain, item) for item in CONGRESSIONAL_OVERSIGHT_DOMAINS):
        return "congressional/federal oversight"
    if domain.endswith(".gov"):
        return "official government"
    if any(base_domain_match(domain, item) for item in RESEARCH_INSTITUTION_DOMAINS):
        return "university/research institution"
    if domain.endswith(".edu"):
        return "university/research institution"
    if any(base_domain_match(domain, item) for item in ESTABLISHED_THINK_TANK_DOMAINS):
        return "established think tank"
    return "unknown"


def infer_file_type(url: str, content_type: str = "", google_file_format: str = "") -> str:
    lowered_url = urllib.parse.urlparse(url).path.lower()
    lowered_content_type = content_type.lower()
    lowered_google_format = google_file_format.lower()
    if lowered_url.endswith(".pdf") or "pdf" in lowered_content_type or "pdf" in lowered_google_format:
        return "PDF"
    if lowered_url.endswith((".doc", ".docx")):
        return "Word document"
    if lowered_url.endswith((".xls", ".xlsx", ".csv")):
        return "Spreadsheet/data file"
    if "html" in lowered_content_type:
        return "HTML"
    return "HTML" if not lowered_url.endswith(".pdf") else "PDF"


def slug_to_title(url: str) -> str:
    parsed = urllib.parse.urlparse(url)
    domain = domain_from_url(url)
    path = urllib.parse.unquote(parsed.path).strip("/")
    if not path:
        return source_organization_for_domain(domain)
    slug = path.split("/")[-1]
    slug = re.sub(r"\.(pdf|html?|aspx?|docx?|xlsx?)$", "", slug, flags=re.IGNORECASE)
    if re.fullmatch(r"gao-\d{2}-\d+", slug, flags=re.IGNORECASE):
        return slug.upper()
    if base_domain_match(domain, "cbo.gov") and re.fullmatch(r"\d+", slug):
        return f"CBO Publication {slug}"
    if base_domain_match(domain, "congress.gov") and re.fullmatch(
        r"R\d+", slug, flags=re.IGNORECASE
    ):
        return f"CRS Report {slug.upper()}"

    slug = re.sub(r"[-_]+", " ", slug)
    slug = re.sub(r"\s+", " ", slug).strip()
    if not slug:
        return source_organization_for_domain(domain)
    return polish_title(slug.title())


def polish_title(title: str) -> str:
    replacements = {
        " Ai ": " AI ",
        " Fy": " FY",
        " Cbo ": " CBO ",
        " Gao ": " GAO ",
        " Crs ": " CRS ",
        " Kff ": " KFF ",
    }
    polished = f" {title} "
    for old, new in replacements.items():
        polished = polished.replace(old, new)
    return polished.strip()


def is_generic_title(title: str, document: PolicyDocument) -> bool:
    lowered = clean_text(title).lower()
    generic_values = {
        "",
        document.domain.lower(),
        document.source_organization.lower(),
        "u.s. gao",
        "gao",
        "congress.gov",
        "congressional research service",
    }
    return lowered in generic_values or len(lowered) < 8


def normalize_url_key(url: str) -> str:
    parsed = urllib.parse.urlparse(url)
    normalized_path = parsed.path.rstrip("/")
    return urllib.parse.urlunparse(
        (parsed.scheme.lower(), parsed.netloc.lower(), normalized_path, "", "", "")
    )


def extract_year(text: str) -> str:
    current_year = datetime.now().year
    years = re.findall(r"\b(19[7-9]\d|20\d{2})\b", text)
    valid_years = [int(year) for year in years if 1970 <= int(year) <= current_year + 1]
    if not valid_years:
        gao_match = re.search(r"\bgao[-_](\d{2})[-_]\d+", text, flags=re.IGNORECASE)
        if gao_match:
            inferred_year = 2000 + int(gao_match.group(1))
            if 2000 <= inferred_year <= current_year + 1:
                return str(inferred_year)
        return ""
    return str(max(valid_years))


def detect_policy_topic(text: str, requested_topic: str = "") -> str:
    if requested_topic:
        requested_words = [word for word in re.split(r"\W+", requested_topic.lower()) if word]
        if any(word in text.lower() for word in requested_words):
            return requested_topic

    scored_topics: List[Tuple[int, str]] = []
    for topic, keywords in POLICY_TOPIC_KEYWORDS.items():
        hits = keyword_hits(text, keywords)
        if hits:
            scored_topics.append((len(hits), topic))
    if not scored_topics:
        return requested_topic or "general public policy"
    scored_topics.sort(reverse=True)
    return scored_topics[0][1]


def infer_document_type(domain: str, title: str, url: str, file_type: str) -> str:
    text = f"{title} {url}".lower()
    if "policy brief" in text or "brief" in text:
        return "policy brief"
    if base_domain_match(domain, "congress.gov") or "bill" in text or "law" in text:
        return "legal/government document"
    if domain.endswith(".gov") or any(
        base_domain_match(domain, item) for item in CONGRESSIONAL_OVERSIGHT_DOMAINS
    ):
        return "government report"
    if any(base_domain_match(domain, item) for item in ESTABLISHED_THINK_TANK_DOMAINS):
        return "think tank report"
    if "report" in text or file_type == "PDF":
        return "research report"
    if domain.endswith(".edu") or any(
        base_domain_match(domain, item) for item in RESEARCH_INSTITUTION_DOMAINS
    ):
        return "research report"
    return "research report"


def build_relevance_note(document: PolicyDocument, requested_topic: str = "") -> str:
    topic = requested_topic or document.policy_topic
    if topic and topic != "general public policy":
        return (
            f"Potentially useful for the AI Playbook because it connects {topic} "
            f"to an authoritative {document.document_type} from {document.source_organization}."
        )
    return (
        f"Potentially useful for the AI Playbook because it provides public policy "
        f"context from {document.source_organization}."
    )


def build_queries(topic: str, quick: bool) -> List[str]:
    queries = GOOGLE_POLICY_QUERIES[:]
    if topic:
        queries = [f"{query} {topic}" for query in queries]
    if quick:
        return queries[:3]
    return queries


def http_request(
    url: str,
    *,
    method: str = "GET",
    max_bytes: int = 500_000,
    max_retries: int = 2,
    sleep_seconds: float = DEFAULT_SLEEP_SECONDS,
    timeout_seconds: float = 10.0,
) -> HttpResult:
    """Request a URL with a clear User-Agent, rate limiting, and simple retries."""
    headers = {
        "User-Agent": USER_AGENT,
        "Accept": "text/html,application/pdf,application/json,text/plain,*/*;q=0.8",
    }
    retry_statuses = {429, 500, 502, 503, 504}

    for attempt in range(1, max_retries + 1):
        time.sleep(sleep_seconds)
        try:
            if requests is not None:
                response = requests.request(
                    method,
                    url,
                    headers=headers,
                    timeout=timeout_seconds,
                    stream=True,
                    allow_redirects=True,
                )
                body = b""
                if method.upper() != "HEAD":
                    chunks = []
                    total = 0
                    for chunk in response.iter_content(chunk_size=8192):
                        if not chunk:
                            continue
                        remaining = max_bytes - total
                        if remaining <= 0:
                            break
                        chunks.append(chunk[:remaining])
                        total += len(chunk[:remaining])
                    body = b"".join(chunks)

                if response.status_code in retry_statuses and attempt < max_retries:
                    continue
                return HttpResult(
                    status_code=response.status_code,
                    headers={key.lower(): value for key, value in response.headers.items()},
                    body=body,
                    final_url=response.url,
                )

            request = urllib.request.Request(url, headers=headers, method=method)
            with urllib.request.urlopen(request, timeout=timeout_seconds) as response:
                body = b"" if method.upper() == "HEAD" else response.read(max_bytes)
                return HttpResult(
                    status_code=response.getcode(),
                    headers={key.lower(): value for key, value in response.headers.items()},
                    body=body,
                    final_url=response.geturl(),
                )
        except urllib.error.HTTPError as exc:
            if exc.code in retry_statuses and attempt < max_retries:
                continue
            return HttpResult(
                status_code=exc.code,
                headers={key.lower(): value for key, value in exc.headers.items()},
                body=b"",
                final_url=url,
                error=str(exc),
            )
        except Exception as exc:  # Network and requests exceptions vary by platform.
            if attempt < max_retries:
                continue
            return HttpResult(
                status_code=0,
                headers={},
                body=b"",
                final_url=url,
                error=str(exc),
            )

    return HttpResult(status_code=0, headers={}, body=b"", final_url=url, error="Unknown error")


def decode_body(body: bytes, headers: Dict[str, str]) -> str:
    content_type = headers.get("content-type", "")
    charset_match = re.search(r"charset=([\w.-]+)", content_type, flags=re.IGNORECASE)
    charset = charset_match.group(1) if charset_match else "utf-8"
    try:
        return body.decode(charset, errors="replace")
    except LookupError:
        return body.decode("utf-8", errors="replace")


def request_json(
    url: str,
    params: Dict[str, Any],
    *,
    sleep_seconds: float,
    timeout_seconds: float,
    max_retries: int = 2,
) -> Optional[Dict[str, Any]]:
    encoded_params = urllib.parse.urlencode(
        {key: value for key, value in params.items() if value not in (None, "")}
    )
    request_url = f"{url}?{encoded_params}"
    result = http_request(
        request_url,
        max_bytes=1_000_000,
        max_retries=max_retries,
        sleep_seconds=sleep_seconds,
        timeout_seconds=timeout_seconds,
    )
    if result.status_code < 200 or result.status_code >= 400:
        print(f"  Request failed ({result.status_code}) for API call: {result.error}")
        return None
    try:
        return json.loads(decode_body(result.body, result.headers))
    except json.JSONDecodeError as exc:
        print(f"  Could not decode API JSON: {exc}")
        return None


def parse_meta_tags(html_text: str) -> Dict[str, str]:
    metadata: Dict[str, str] = {}
    for match in re.finditer(r"<meta\s+([^>]+)>", html_text, flags=re.IGNORECASE):
        attributes = {
            key.lower(): html.unescape(value)
            for key, value in re.findall(
                r'([a-zA-Z_:.-]+)\s*=\s*["\']([^"\']*)["\']', match.group(1)
            )
        }
        name = attributes.get("name") or attributes.get("property")
        content = attributes.get("content")
        if name and content:
            metadata[name.lower()] = clean_text(content)
    return metadata


def parse_html_title(html_text: str) -> str:
    metadata = parse_meta_tags(html_text)
    for key in ("og:title", "twitter:title", "citation_title", "dc.title"):
        if metadata.get(key):
            return strip_site_suffix(metadata[key])

    title_match = re.search(
        r"<title[^>]*>(.*?)</title>",
        html_text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if title_match:
        return strip_site_suffix(clean_text(title_match.group(1)))
    return ""


def strip_site_suffix(title: str) -> str:
    title = clean_text(title)
    for separator in (" | ", " - "):
        parts = title.split(separator)
        if len(parts) > 1 and len(parts[0]) >= 8:
            return parts[0].strip()
    return title


def parse_html_description(html_text: str) -> str:
    metadata = parse_meta_tags(html_text)
    for key in ("description", "og:description", "twitter:description", "dc.description"):
        if metadata.get(key):
            return metadata[key]
    return ""


def parse_html_year(html_text: str) -> str:
    metadata = parse_meta_tags(html_text)
    year_candidates = []
    for key, value in metadata.items():
        if any(token in key for token in ("date", "published", "created", "modified")):
            year = extract_year(value)
            if year:
                year_candidates.append(year)
    if year_candidates:
        return max(year_candidates)
    return extract_year(html_text[:20_000])


def document_from_url(
    url: str,
    *,
    source_mode: str,
    requested_topic: str = "",
    title: str = "",
    snippet: str = "",
    google_file_format: str = "",
) -> PolicyDocument:
    domain = domain_from_url(url)
    title = clean_text(title) or slug_to_title(url)
    source_organization = source_organization_for_domain(domain)
    file_type = infer_file_type(url, google_file_format=google_file_format)
    combined_text = f"{title} {snippet} {url}"
    publication_year = extract_year(combined_text)
    policy_topic = detect_policy_topic(combined_text, requested_topic)
    document_type = infer_document_type(domain, title, url, file_type)
    authority_level = authority_level_for_domain(domain)

    document = PolicyDocument(
        title=title,
        source_organization=source_organization,
        url=url,
        domain=domain,
        file_type=file_type,
        publication_year=publication_year,
        policy_topic=policy_topic,
        document_type=document_type,
        authority_level=authority_level,
        relevance_note="",
        retrieval_date=date.today().isoformat(),
        source_mode=source_mode,
        snippet=clean_text(snippet),
    )
    document.relevance_note = build_relevance_note(document, requested_topic)
    return document


def enrich_document_from_web(
    document: PolicyDocument,
    robots_cache: RobotsCache,
    *,
    requested_topic: str,
    sleep_seconds: float,
    timeout_seconds: float,
    quick: bool,
) -> PolicyDocument:
    allowed, robots_note = robots_cache.can_fetch(document.url)
    if not allowed:
        document.fetch_status = robots_note
        document.relevance_note += " Content fetching was skipped because robots.txt disallows it."
        return document

    result = http_request(
        document.url,
        max_bytes=350_000 if quick else 750_000,
        max_retries=1 if quick else 2,
        sleep_seconds=sleep_seconds,
        timeout_seconds=timeout_seconds,
    )
    if result.status_code < 200 or result.status_code >= 400:
        document.fetch_status = f"Fetch skipped/failed ({result.status_code}): {result.error}"
        return document

    content_type = result.headers.get("content-type", "")
    document.file_type = infer_file_type(
        result.final_url or document.url,
        content_type=content_type,
    )
    if result.final_url:
        document.url = result.final_url
        document.domain = domain_from_url(result.final_url)

    if "html" not in content_type.lower():
        document.fetch_status = f"Fetched headers/content type only ({document.file_type})"
        return document

    html_text = decode_body(result.body, result.headers)
    fetched_title = parse_html_title(html_text)
    if fetched_title and not is_generic_title(fetched_title, document):
        document.title = fetched_title
    description = parse_html_description(html_text)
    combined_text = f"{document.title} {description} {document.snippet} {document.url}"
    fetched_year = parse_html_year(html_text)
    if fetched_year:
        document.publication_year = fetched_year
    document.policy_topic = detect_policy_topic(combined_text, requested_topic)
    document.document_type = infer_document_type(
        document.domain,
        document.title,
        document.url,
        document.file_type,
    )
    document.source_organization = source_organization_for_domain(document.domain)
    document.authority_level = authority_level_for_domain(document.domain)
    document.relevance_note = build_relevance_note(document, requested_topic)
    document.fetch_status = "Fetched public HTML metadata"
    return document


def search_google_api(args: argparse.Namespace) -> List[PolicyDocument]:
    api_key = os.getenv("GOOGLE_API_KEY", "")
    cse_id = os.getenv("GOOGLE_CSE_ID", "")
    if not api_key or not cse_id:
        return []

    queries = build_queries(args.topic, args.quick)
    results_per_query = 3 if args.quick else 10
    documents: List[PolicyDocument] = []

    print("Google API mode")
    print("Using the official Google Custom Search JSON API.")
    print(f"Queries: {len(queries)}")
    for index, query in enumerate(queries, start=1):
        if len(documents) >= args.max_results:
            break
        remaining = args.max_results - len(documents)
        num_results = min(results_per_query, remaining, 10)
        print(f"[{index}/{len(queries)}] {query}")

        data = request_json(
            GOOGLE_CUSTOM_SEARCH_ENDPOINT,
            {
                "key": api_key,
                "cx": cse_id,
                "q": query,
                "num": num_results,
            },
            sleep_seconds=args.sleep_seconds,
            timeout_seconds=args.timeout_seconds,
            max_retries=args.max_retries,
        )
        if not data:
            record_note(args, f"Google API request returned no usable data for query: {query}")
            continue

        for item in data.get("items", []):
            link = clean_text(item.get("link"))
            if not link:
                record_note(args, f"Skipped Google API item without a URL for query: {query}")
                continue
            documents.append(
                document_from_url(
                    link,
                    source_mode="Google API",
                    requested_topic=args.topic,
                    title=clean_text(item.get("title")),
                    snippet=clean_text(item.get("snippet")),
                    google_file_format=clean_text(item.get("fileFormat")),
                )
            )
            if len(documents) >= args.max_results:
                break
        print(f"  Collected so far: {len(documents)}")

    return documents


def load_seed_urls(seed_file: str) -> List[str]:
    if not os.path.exists(seed_file):
        raise FileNotFoundError(
            f"Seed file not found: {seed_file}. Create one URL per line or set "
            "GOOGLE_API_KEY and GOOGLE_CSE_ID for Google API mode."
        )

    urls = []
    with open(seed_file, "r", encoding="utf-8") as file:
        for line in file:
            stripped = line.strip()
            if not stripped or stripped.startswith("#"):
                continue
            urls.append(stripped)
    return urls


def search_manual_seed_file(args: argparse.Namespace) -> List[PolicyDocument]:
    urls = load_seed_urls(args.seed_file)
    limit = min(args.max_results, len(urls))
    print("Manual URL mode")
    print("Using manual URL mode because Google API credentials were not found.")
    print(f"Seed file: {args.seed_file}")
    print(f"Seed URLs loaded: {len(urls)}")
    for skipped_url in urls[limit:]:
        record_note(args, f"Skipped due max-results limit: {skipped_url}")

    documents = [
        document_from_url(
            url,
            source_mode="Manual URL",
            requested_topic=args.topic,
        )
        for url in urls[:limit]
    ]
    return documents


def deduplicate_documents(documents: Iterable[PolicyDocument]) -> List[PolicyDocument]:
    seen: Dict[str, PolicyDocument] = {}
    for document in documents:
        key = normalize_url_key(document.url)
        if key and key not in seen:
            seen[key] = document
    return list(seen.values())


def enrich_documents(documents: List[PolicyDocument], args: argparse.Namespace) -> List[PolicyDocument]:
    if args.quick:
        print("Quick mode: skipping robots.txt/page/PDF fetching and using URL-level metadata only.")
        for document in documents:
            document.fetch_status = "Quick mode URL-level metadata only"
        return documents

    fetch_limit = min(len(documents), 3 if args.quick else args.max_results)
    if fetch_limit <= 0:
        return documents

    print(f"Fetching public page metadata for up to {fetch_limit} document(s).")
    robots_cache = RobotsCache(
        sleep_seconds=args.sleep_seconds,
        timeout_seconds=args.timeout_seconds,
        quick=args.quick,
    )
    enriched = []
    for index, document in enumerate(documents, start=1):
        if index <= fetch_limit:
            print(f"  [{index}/{fetch_limit}] {document.domain}")
            enriched.append(
                enrich_document_from_web(
                    document,
                    robots_cache,
                    requested_topic=args.topic,
                    sleep_seconds=args.sleep_seconds,
                    timeout_seconds=args.timeout_seconds,
                    quick=args.quick,
                )
            )
        else:
            document.fetch_status = "Not fetched due demo/result limit"
            enriched.append(document)
    return enriched


def write_csv(documents: List[PolicyDocument], path: str) -> None:
    with open(path, "w", newline="", encoding="utf-8") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=CSV_COLUMNS)
        writer.writeheader()
        for document in documents:
            writer.writerow({column: getattr(document, column) for column in CSV_COLUMNS})


def write_docx_report(
    documents: List[PolicyDocument],
    path: str,
    args: argparse.Namespace,
    mode: str,
) -> None:
    """Write a Word report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to write the Word report. "
            "Install it with: pip install -r requirements.txt"
        ) from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    document.add_heading("AI Playbook Public Policy Document Finder Report", level=1)
    document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    document.add_heading("Purpose", level=2)
    document.add_paragraph(
        "This report organizes authoritative public policy documents that can "
        "support an AI Playbook workflow. It is a starting point for source "
        "selection, prompt grounding, and manual review rather than a substitute "
        "for reading the documents."
    )

    document.add_heading("Search Strategy", level=2)
    search_items = [
        f"Source mode used: {display_source_mode(mode)}",
        f"Google API endpoint: {GOOGLE_CUSTOM_SEARCH_ENDPOINT}",
        f"Seed file: {args.seed_file}",
        f"Maximum documents requested: {args.max_results}",
        f"Quick demo mode: {'yes' if args.quick else 'no'}",
        f"Also wrote latest copies: {'yes' if args.also_write_latest else 'no'}",
        f"Optional topic: {args.topic or 'none'}",
        (
            "Responsible crawling: Google result pages were not scraped; "
            "robots.txt was checked before content fetching in normal mode; "
            "paywalls and logins were not bypassed."
        ),
    ]
    for item in search_items:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Summary Counts", level=2)
    document.add_paragraph(f"Total candidate documents: {len(documents)}")
    document.add_paragraph(f"Retrieval date: {date.today().isoformat()}")
    for authority_level, count in count_by(documents, "authority_level").items():
        document.add_paragraph(
            f"{authority_level}: {count}",
            style="List Bullet",
        )

    document.add_heading("Candidate Documents by Authority Level", level=2)
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = [
        "Title",
        "Source organization",
        "Authority level",
        "Document type",
        "File type",
        "URL",
    ]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    sorted_documents = sorted(
        documents,
        key=lambda item: (item.authority_level, item.source_organization, item.title),
    )
    for candidate in sorted_documents:
        cells = table.add_row().cells
        cells[0].text = candidate.title
        cells[1].text = candidate.source_organization
        cells[2].text = candidate.authority_level
        cells[3].text = candidate.document_type
        cells[4].text = candidate.file_type
        cells[5].text = candidate.url

    grouped: Dict[str, List[PolicyDocument]] = {}
    for candidate in documents:
        grouped.setdefault(candidate.authority_level or "unknown", []).append(candidate)

    authority_order = [
        "official government",
        "congressional/federal oversight",
        "university/research institution",
        "established think tank",
        "unknown",
    ]
    ordered_levels = authority_order + [
        level for level in grouped if level not in authority_order
    ]

    for authority_level in ordered_levels:
        candidates = grouped.get(authority_level, [])
        if not candidates:
            continue
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Authority level: {authority_level}")
        run.bold = True
        for candidate in candidates:
            document.add_heading(candidate.title, level=3)
            details = [
                f"Source organization: {candidate.source_organization}",
                f"URL: {candidate.url}",
                f"Domain: {candidate.domain}",
                f"File type: {candidate.file_type}",
                f"Publication year: {candidate.publication_year or 'Unknown'}",
                f"Policy topic: {candidate.policy_topic}",
                f"Document type: {candidate.document_type}",
                f"Relevance note: {candidate.relevance_note}",
                f"Source mode: {candidate.source_mode}",
            ]
            for detail in details:
                document.add_paragraph(detail, style="List Bullet")

    document.add_heading("How This Supports the AI Playbook", level=2)
    document.add_paragraph(
        "The CSV gives a structured source list that can feed an AI Playbook "
        "workflow: documents can be assigned to policy tasks, checked for "
        "authority and currency, and used as grounding material before AI "
        "summaries or recommendations are produced."
    )

    document.add_heading("Manual Verification Checklist", level=2)
    checklist_items = [
        "Confirm the document is publicly accessible.",
        "Confirm the source is authoritative and relevant.",
        "Confirm the document is current enough for the policy issue.",
        "Confirm the script did not use restricted or paywalled content.",
        "Confirm the AI Playbook user reads the source before relying on AI-generated summaries.",
    ]
    for item in checklist_items:
        document.add_paragraph(item, style="List Bullet")

    document.save(path)


def markdown_escape(text: str) -> str:
    return clean_text(text).replace("|", "\\|")


def count_by(documents: List[PolicyDocument], field_name: str) -> Dict[str, int]:
    counts: Dict[str, int] = {}
    for document in documents:
        value = getattr(document, field_name) or "Unknown"
        counts[value] = counts.get(value, 0) + 1
    return dict(sorted(counts.items(), key=lambda item: (-item[1], item[0])))


def create_run_paths() -> RunPaths:
    """Create a timestamped run folder and return all run-specific paths."""
    os.makedirs(RUNS_DIR, exist_ok=True)
    base_run_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    counter = 1
    while True:
        run_id = base_run_id if counter == 1 else f"{base_run_id}_{counter:02d}"
        run_folder = os.path.join(RUNS_DIR, f"run_{run_id}")
        try:
            os.makedirs(run_folder, exist_ok=False)
            break
        except FileExistsError:
            counter += 1

    return RunPaths(
        run_id=run_id,
        run_folder=run_folder,
        csv_path=os.path.join(run_folder, f"policy_document_candidates_{run_id}.csv"),
        report_path=os.path.join(run_folder, f"policy_document_report_{run_id}.md"),
        docx_path=os.path.join(run_folder, f"policy_document_report_{run_id}.docx"),
        log_path=os.path.join(run_folder, f"run_log_{run_id}.txt"),
    )


def display_path(path: str, trailing_slash: bool = False) -> str:
    shown = path.replace(os.sep, "/")
    if trailing_slash and not shown.endswith("/"):
        shown += "/"
    return shown


def record_note(args: argparse.Namespace, note: str) -> None:
    if not hasattr(args, "run_notes"):
        args.run_notes = []
    args.run_notes.append(note)


def display_source_mode(mode: str) -> str:
    if mode.endswith("mode"):
        return mode
    return f"{mode} mode"


def collect_skipped_or_error_notes(
    documents: List[PolicyDocument],
    args: argparse.Namespace,
) -> List[str]:
    notes = list(getattr(args, "run_notes", []))
    for document in documents:
        status = document.fetch_status
        if not status or status in {
            "Not fetched",
            "Fetched public HTML metadata",
            "Fetched headers/content type only (PDF)",
        }:
            continue
        notes.append(f"{document.url}: {status}")
    if args.quick:
        notes.append("Quick mode skipped deep content fetching for all candidates.")
    return notes


def write_run_log(
    documents: List[PolicyDocument],
    path: str,
    args: argparse.Namespace,
    mode: str,
    run_paths: RunPaths,
    latest_paths: List[str],
) -> None:
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    skipped_or_error_notes = collect_skipped_or_error_notes(documents, args)
    output_paths = [
        run_paths.csv_path,
        run_paths.report_path,
        run_paths.docx_path,
        run_paths.log_path,
        RUN_INDEX_CSV,
    ]
    output_paths.extend(latest_paths)

    lines = [
        "AI Playbook Public Policy Document Finder Run Log",
        "",
        f"Run ID: {run_paths.run_id}",
        f"Date and time: {generated_at}",
        f"Command line arguments used: {' '.join(sys.argv[1:]) or '(none)'}",
        f"Mode used: {'Quick mode' if args.quick else 'Normal mode'}",
        f"Source mode: {display_source_mode(mode)}",
        f"Seed file used: {args.seed_file}",
        f"Topic used: {args.topic or 'none'}",
        f"Max results requested: {args.max_results}",
        f"Candidate documents collected: {len(documents)}",
        "",
        "Output file paths:",
    ]
    lines.extend(f"- {display_path(path_item)}" for path_item in output_paths)
    lines.extend(["", "Skipped URLs or errors:"])
    if skipped_or_error_notes:
        lines.extend(f"- {note}" for note in skipped_or_error_notes)
    else:
        lines.append("- None recorded.")
    lines.extend(
        [
            "",
            "Manual verification reminder:",
            "- Manual verification required before relying on any source.",
            "",
        ]
    )

    with open(path, "w", encoding="utf-8") as log_file:
        log_file.write("\n".join(lines))


def append_run_index(
    documents: List[PolicyDocument],
    args: argparse.Namespace,
    mode: str,
    run_paths: RunPaths,
) -> None:
    """Append a single successful run to the persistent run index."""
    os.makedirs(RUNS_DIR, exist_ok=True)
    ensure_run_index_schema()
    file_exists = os.path.exists(RUN_INDEX_CSV)
    row = {
        "run_id": run_paths.run_id,
        "datetime": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "quick_mode": "yes" if args.quick else "no",
        "source_mode": display_source_mode(mode),
        "seed_file": args.seed_file,
        "topic": args.topic,
        "max_results_requested": str(args.max_results),
        "documents_collected": str(len(documents)),
        "run_folder": display_path(run_paths.run_folder),
        "csv_output_path": display_path(run_paths.csv_path),
        "markdown_report_path": display_path(run_paths.report_path),
        "word_report_path": display_path(run_paths.docx_path),
        "log_path": display_path(run_paths.log_path),
        "also_write_latest": "yes" if args.also_write_latest else "no",
    }
    with open(RUN_INDEX_CSV, "a", newline="", encoding="utf-8") as index_file:
        writer = csv.DictWriter(index_file, fieldnames=RUN_INDEX_COLUMNS)
        if not file_exists or os.path.getsize(RUN_INDEX_CSV) == 0:
            writer.writeheader()
        writer.writerow(row)


def ensure_run_index_schema() -> None:
    """Upgrade older run_index.csv headers while preserving previous rows."""
    if not os.path.exists(RUN_INDEX_CSV) or os.path.getsize(RUN_INDEX_CSV) == 0:
        return

    with open(RUN_INDEX_CSV, "r", newline="", encoding="utf-8") as index_file:
        reader = csv.DictReader(index_file)
        existing_columns = reader.fieldnames or []
        rows = list(reader)

    if existing_columns == RUN_INDEX_COLUMNS:
        return

    legacy_copy = os.path.join(
        RUNS_DIR,
        f"run_index_legacy_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}.csv",
    )
    os.replace(RUN_INDEX_CSV, legacy_copy)
    with open(RUN_INDEX_CSV, "w", newline="", encoding="utf-8") as index_file:
        writer = csv.DictWriter(index_file, fieldnames=RUN_INDEX_COLUMNS)
        writer.writeheader()
        for row in rows:
            writer.writerow({column: row.get(column, "") for column in RUN_INDEX_COLUMNS})


def write_report(documents: List[PolicyDocument], path: str, args: argparse.Namespace, mode: str) -> None:
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    grouped: Dict[str, List[PolicyDocument]] = {}
    for document in documents:
        grouped.setdefault(document.authority_level or "unknown", []).append(document)

    lines = [
        "# AI Playbook Public Policy Document Finder Report",
        "",
        f"Generated: {generated_at}",
        "",
        "## Purpose",
        "",
        (
            "This report organizes authoritative public policy documents that can "
            "support an AI Playbook workflow. It is a starting point for source "
            "selection, prompt grounding, and manual review rather than a substitute "
            "for reading the documents."
        ),
        "",
        "## Search Strategy",
        "",
        f"- Source mode used: {display_source_mode(mode)}",
        f"- Google API endpoint: {GOOGLE_CUSTOM_SEARCH_ENDPOINT}",
        f"- Seed file: {args.seed_file}",
        f"- Maximum documents requested: {args.max_results}",
        f"- Quick demo mode: {'yes' if args.quick else 'no'}",
        f"- Also wrote latest copies: {'yes' if args.also_write_latest else 'no'}",
        f"- Optional topic: {args.topic or 'none'}",
        f"- Responsible crawling: Google result pages were not scraped; robots.txt was checked before content fetching; paywalls and logins were not bypassed.",
        "",
        "## Summary Counts",
        "",
        f"- Total candidate documents: {len(documents)}",
        f"- Retrieval date: {date.today().isoformat()}",
        "",
        "### By Authority Level",
        "",
        "| Authority level | Count |",
        "|---|---:|",
    ]

    for authority_level, count in count_by(documents, "authority_level").items():
        lines.append(f"| {markdown_escape(authority_level)} | {count} |")

    lines.extend(["", "### By File Type", "", "| File type | Count |", "|---|---:|"])
    for file_type, count in count_by(documents, "file_type").items():
        lines.append(f"| {markdown_escape(file_type)} | {count} |")

    lines.extend(["", "## Candidate Documents by Authority Level", ""])
    authority_order = [
        "official government",
        "congressional/federal oversight",
        "university/research institution",
        "established think tank",
        "unknown",
    ]
    ordered_levels = authority_order + [
        level for level in grouped if level not in authority_order
    ]

    for authority_level in ordered_levels:
        candidates = grouped.get(authority_level, [])
        if not candidates:
            continue
        lines.extend([f"### {authority_level}", ""])
        for index, document in enumerate(candidates, start=1):
            lines.extend(
                [
                    f"#### {index}. {document.title}",
                    "",
                    f"- Source organization: {document.source_organization}",
                    f"- URL: {document.url}",
                    f"- Domain: {document.domain}",
                    f"- File type: {document.file_type}",
                    f"- Publication year: {document.publication_year or 'Unknown'}",
                    f"- Policy topic: {document.policy_topic}",
                    f"- Document type: {document.document_type}",
                    f"- Relevance note: {document.relevance_note}",
                    f"- Source mode: {document.source_mode}",
                    "",
                ]
            )

    lines.extend(
        [
            "## How This Supports the AI Playbook",
            "",
            (
                "The CSV gives a structured source list that can feed an AI Playbook "
                "workflow: documents can be assigned to policy tasks, checked for "
                "authority and currency, and used as grounding material before AI "
                "summaries or recommendations are produced."
            ),
            "",
            "## Manual Verification Checklist",
            "",
            "- Confirm the document is publicly accessible.",
            "- Confirm the source is authoritative and relevant.",
            "- Confirm the document is current enough for the policy issue.",
            "- Confirm the script did not use restricted or paywalled content.",
            "- Confirm the AI Playbook user reads the source before relying on AI-generated summaries.",
            "",
        ]
    )

    with open(path, "w", encoding="utf-8") as md_file:
        md_file.write("\n".join(lines))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Find and organize authoritative public policy document candidates "
            "for an AI Playbook workflow."
        )
    )
    parser.add_argument(
        "--quick",
        action="store_true",
        help="One-minute demo mode with fewer queries/results and limited fetching.",
    )
    parser.add_argument(
        "--seed-file",
        default=DEFAULT_SEED_FILE,
        help=f"Manual URL seed file. Default: {DEFAULT_SEED_FILE}",
    )
    parser.add_argument(
        "--max-results",
        type=int,
        default=None,
        help=(
            "Maximum number of documents to collect. "
            f"Default: {DEFAULT_MAX_RESULTS}; quick-mode default: {QUICK_MAX_RESULTS}"
        ),
    )
    parser.add_argument(
        "--topic",
        default="",
        help='Optional topic keyword to include in Google API queries, e.g. "housing policy".',
    )
    parser.add_argument(
        "--sleep-seconds",
        type=float,
        default=DEFAULT_SLEEP_SECONDS,
        help="Seconds to sleep before requests for rate limiting.",
    )
    parser.add_argument(
        "--also-write-latest",
        action="store_true",
        help=(
            "Also write policy_document_candidates.csv and "
            "policy_document_report.md as overwriteable latest copies."
        ),
    )
    return parser.parse_args()


def validate_args(args: argparse.Namespace) -> None:
    if args.max_results is not None and args.max_results < 1:
        raise ValueError("--max-results must be at least 1")
    if args.sleep_seconds < 0:
        raise ValueError("--sleep-seconds cannot be negative")
    args.topic = clean_text(args.topic)
    if args.max_results is None:
        args.max_results = QUICK_MAX_RESULTS if args.quick else DEFAULT_MAX_RESULTS

    if args.quick:
        args.timeout_seconds = QUICK_TIMEOUT_SECONDS
        args.max_retries = 1
        args.sleep_seconds = min(args.sleep_seconds, QUICK_SLEEP_SECONDS)
    else:
        args.timeout_seconds = DEFAULT_TIMEOUT_SECONDS
        args.max_retries = 2
    args.run_notes = []


def choose_mode_and_collect(args: argparse.Namespace) -> Tuple[List[PolicyDocument], str]:
    if args.quick:
        print("Quick mode enabled")

    has_google_credentials = bool(os.getenv("GOOGLE_API_KEY") and os.getenv("GOOGLE_CSE_ID"))
    if has_google_credentials:
        documents = search_google_api(args)
        mode = "Google API"
    else:
        documents = search_manual_seed_file(args)
        mode = "Manual URL"

    documents = deduplicate_documents(documents)
    documents = documents[: args.max_results]
    documents = enrich_documents(documents, args)
    return documents, mode


def print_summary(
    documents: List[PolicyDocument],
    mode: str,
    run_paths: RunPaths,
    latest_paths: List[str],
    args: argparse.Namespace,
) -> None:
    print("\nRun complete.")
    print("Quick mode enabled" if args.quick else "Normal mode enabled")
    print(display_source_mode(mode))
    print(f"Collected {len(documents)} candidate documents")
    print(f"Created run folder: {display_path(run_paths.run_folder, trailing_slash=True)}")
    created_files = [
        run_paths.csv_path,
        run_paths.report_path,
        run_paths.docx_path,
        run_paths.log_path,
    ]
    created_files.extend(latest_paths)
    print(f"Created output files: {', '.join(display_path(path) for path in created_files)}")
    print(f"Updated run index: {display_path(RUN_INDEX_CSV)}")
    print("Manual verification required before relying on any source.")


def main() -> int:
    try:
        args = parse_args()
        validate_args(args)
        run_paths = create_run_paths()
        documents, mode = choose_mode_and_collect(args)
        write_csv(documents, run_paths.csv_path)
        write_report(documents, run_paths.report_path, args, mode)
        write_docx_report(documents, run_paths.docx_path, args, mode)

        latest_paths: List[str] = []
        if args.also_write_latest:
            write_csv(documents, LATEST_OUTPUT_CSV)
            write_report(documents, LATEST_OUTPUT_MD, args, mode)
            write_docx_report(documents, LATEST_OUTPUT_DOCX, args, mode)
            latest_paths = [LATEST_OUTPUT_CSV, LATEST_OUTPUT_MD, LATEST_OUTPUT_DOCX]

        write_run_log(documents, run_paths.log_path, args, mode, run_paths, latest_paths)
        append_run_index(documents, args, mode, run_paths)
        print_summary(documents, mode, run_paths, latest_paths, args)
        return 0
    except KeyboardInterrupt:
        print("\nRun cancelled by user.")
        return 130
    except Exception as exc:
        print(f"Error: {exc}")
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
